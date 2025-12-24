import subprocess
from datetime import datetime
from docx import Document

print("Choose mode:")
print("1. Blog Ideas")
print("2. Full SEO Blog")
print("3. Rewrite Content")

choice = input("Enter choice (1/2/3): ")

topic = input("\nEnter topic or content:\n")
keywords = input("\nEnter SEO keywords (comma separated): ")
tone = input("\nChoose tone (professional / casual / persuasive): ")

if choice == "1":
    prompt = f"""
    Generate 10 high-quality blog ideas for startups.

    Topic: {topic}
    SEO Keywords: {keywords}
    Tone: {tone}
    """

elif choice == "2":
    prompt = f"""
    Write a 700-word SEO-optimized blog for a startup.

    Topic: {topic}
    SEO Keywords: {keywords}
    Tone: {tone}

    Include:
    - SEO-friendly headings
    - Introduction and conclusion
    - Natural keyword usage
    """

elif choice == "3":
    prompt = f"""
    Rewrite the following content to be SEO-optimized and engaging
    for startup audiences.

    Content:
    {topic}

    SEO Keywords: {keywords}
    Tone: {tone}
    """

else:
    print("Invalid choice")
    exit()

result = subprocess.run(
    ["ollama", "run", "llama3:8b", prompt],
    capture_output=True,
    text=True
)

output = result.stdout

# 🔹 File names
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
txt_file = f"generated_content_{timestamp}.txt"
docx_file = f"generated_content_{timestamp}.docx"

# 🔹 Save TXT
with open(txt_file, "w") as f:
    f.write(output)

# 🔹 Save DOCX
doc = Document()
doc.add_heading("Generated Content", level=1)

for line in output.split("\n"):
    doc.add_paragraph(line)

doc.save(docx_file)

print("\n--- GENERATED CONTENT ---\n")
print(output)
print(f"\n✅ Saved files:")
print(f"- {txt_file}")
print(f"- {docx_file}")

