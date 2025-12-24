import streamlit as st
from datetime import datetime
from docx import Document
import os
import requests

# -------------------- Page Config --------------------
st.set_page_config(
    page_title="WriteWise",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# -------------------- Custom Styling --------------------
st.markdown(
    """
    <style>
        .block-container {
            padding-top: 2rem;
            padding-bottom: 3rem;
            max-width: 760px;
        }
        h1 {
            margin-bottom: 0.2rem;
        }
        .subtitle {
            color: #6b7280;
            font-size: 0.95rem;
            margin-bottom: 1.2rem;
        }
        .section-label {
            font-weight: 600;
            margin-top: 1.5rem;
            margin-bottom: 0.3rem;
        }
        .helper-text {
            font-size: 0.85rem;
            color: #6b7280;
            margin-bottom: 0.4rem;
        }
    </style>
    """,
    unsafe_allow_html=True
)

# -------------------- Header --------------------
st.title("WriteWise")
st.markdown(
    '<div class="subtitle">Thoughtful writing for startups — clear, original, and ready to publish.</div>',
    unsafe_allow_html=True
)

st.markdown(
    "> WriteWise helps you think through ideas, not just fill pages."
)

# -------------------- Setup --------------------
os.makedirs("outputs", exist_ok=True)

# -------------------- Input Section --------------------
st.markdown('<div class="section-label">What would you like to work on?</div>', unsafe_allow_html=True)
mode = st.selectbox(
    "",
    [
        "Brainstorm blog ideas",
        "Write a complete blog",
        "Refine existing content"
    ]
)

st.markdown('<div class="section-label">Your topic or draft</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="helper-text">Start with a rough idea, headline, or unfinished paragraph.</div>',
    unsafe_allow_html=True
)
topic = st.text_area(
    "",
    placeholder="e.g. How early-stage startups can use AI without overengineering",
    height=140
)

st.markdown('<div class="section-label">Optional keywords</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="helper-text">Only add these if you care about search visibility.</div>',
    unsafe_allow_html=True
)
keywords = st.text_input(
    "",
    placeholder="startup growth, AI tools, productivity"
)

st.markdown('<div class="section-label">Writing style</div>', unsafe_allow_html=True)
tone = st.selectbox(
    "",
    ["Clear & professional", "Conversational", "Persuasive"]
)

st.divider()

# -------------------- Groq API Call --------------------
def call_groq(prompt: str) -> str:
    api_key = st.secrets.get("GROQ_API_KEY")

    if not api_key:
        st.error("API configuration missing. Please try again later.")
        st.stop()

    url = "https://api.groq.com/openai/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "llama-3.1-8b-instant",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You write like a thoughtful human editor. "
                    "Your tone is natural, concise, and clear. "
                    "Avoid generic AI phrases, buzzwords, and filler. "
                    "Write for real startup founders and teams."
                )
            },
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.5,
        "max_tokens": 600
    }

    response = requests.post(url, headers=headers, json=payload, timeout=30)

    if response.status_code != 200:
        st.error("Something went wrong while preparing your draft.")
        st.code(response.text)
        st.stop()

    return response.json()["choices"][0]["message"]["content"]

# -------------------- Action --------------------
st.markdown('<div class="section-label">Ready?</div>', unsafe_allow_html=True)

if st.button("Create draft"):
    if not topic:
        st.warning("Add a topic or draft to continue.")
    else:
        if mode == "Brainstorm blog ideas":
            prompt = f"""
            Suggest 8–10 thoughtful blog ideas for a startup audience.

            Topic:
            {topic}

            Optional keywords:
            {keywords}

            Writing style:
            {tone}

            The ideas should feel practical, specific, and worth reading.
            """

        elif mode == "Write a complete blog":
            prompt = f"""
            Write a clear, well-structured blog post (400–500 words)
            aimed at startup founders or small teams.

            Topic:
            {topic}

            Optional keywords:
            {keywords}

            Writing style:
            {tone}

            Include an introduction, section headings, and a short conclusion.
            The tone should feel human and grounded.
            """

        else:
            prompt = f"""
            Edit and refine the following draft.

            Improve clarity, flow, and tone without sounding artificial.
            Keep the voice human, direct, and confident.

            Draft:
            {topic}

            Optional keywords:
            {keywords}

            Writing style:
            {tone}
            """

        with st.spinner("Putting together a thoughtful draft…"):
            output = call_groq(prompt)

        st.success("Draft ready.")

        st.text_area(
            "Your draft",
            output,
            height=280
        )

        # -------------------- Save Files --------------------
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        txt_path = f"outputs/draft_{timestamp}.txt"
        docx_path = f"outputs/draft_{timestamp}.docx"

        with open(txt_path, "w") as f:
            f.write(output)

        doc = Document()
        doc.add_heading("Draft", level=1)
        for line in output.split("\n"):
            doc.add_paragraph(line)
        doc.save(docx_path)

        st.download_button("Download as text", output, file_name=f"draft_{timestamp}.txt")
        with open(docx_path, "rb") as f:
            st.download_button("Download as Word file", f, file_name=f"draft_{timestamp}.docx")

